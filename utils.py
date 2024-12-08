import os
import subprocess
import sys
from PyPDF2 import PdfReader
from docx import Document
from datetime import datetime
from docx import Document
from langchain_chroma import Chroma
from langchain.prompts import ChatPromptTemplate
from langchain_ollama import OllamaLLM

sys.path.append(os.path.join(os.path.dirname(__file__), 'rag-system'))
from get_embeddings import get_embeddings_function

# File Handling Utilities
def create_directory_if_not_exists(directory_path):
    """Ensure a directory exists."""
    os.makedirs(directory_path, exist_ok=True)


def delete_all_files_in_directory(directory_path):
    """Delete all files in the specified directory."""
    if os.path.exists(directory_path):
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            if os.path.isfile(file_path):
                os.unlink(file_path)


def allowed_file(filename, allowed_extensions):
    """Check if the file type is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions


def count_characters(file_path, file_type):
    """Count the number of characters in a file based on its type."""
    if file_type == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return len(f.read())
    elif file_type == 'pdf':
        reader = PdfReader(file_path)
        return sum(len(page.extract_text()) for page in reader.pages)
    elif file_type in {'doc', 'docx'}:
        doc = Document(file_path)
        return sum(len(paragraph.text) for paragraph in doc.paragraphs)
    return 0


# Validation Utilities
def validate_questions(form_data, question_types):
    """Validate question input data from the form."""
    errors = []
    question_data = []
    for i, question_type in enumerate(question_types, start=1):
        try:
            easy_count = int(form_data.get(f'easy_{i}', 0))
            medium_count = int(form_data.get(f'medium_{i}', 0))
            difficult_count = int(form_data.get(f'difficult_{i}', 0))

            if not (0 <= easy_count <= 10):
                errors.append(f"Easy count for '{question_type}' must be between 0 and 10.")
            if not (0 <= medium_count <= 10):
                errors.append(f"Medium count for '{question_type}' must be between 0 and 10.")
            if not (0 <= difficult_count <= 10):
                errors.append(f"Difficult count for '{question_type}' must be between 0 and 10.")

            question_data.append({
                'question_type': question_type,
                'easy': easy_count,
                'medium': medium_count,
                'difficult': difficult_count
            })
        except ValueError:
            errors.append(f"Invalid input for '{question_type}'. Please enter numeric values.")

    return question_data, errors


def query_rag(query_text, chroma_path, prompt_template):
    """Query the RAG system and retrieve results."""
    # Initialize the Chroma database with the embedding function
    embedding_function = get_embeddings_function()
    db = None

    try:
        db = Chroma(persist_directory=chroma_path, embedding_function=embedding_function)
        results = db.similarity_search_with_score(query_text, k=5)

        if not results:
            return "No relevant context found."

        # Prepare the context text for the model
        context_text = "\n\n---\n\n".join([doc.page_content for doc, _ in results])
        prompt = ChatPromptTemplate.from_template(prompt_template).format(
            context=context_text,
            question=query_text
        )

        # Query the model
        model = OllamaLLM(model="llama3.2:3b")
        return model.invoke(prompt)

    except Exception as e:
        print(f"Error querying RAG system: {e}")
        return "An error occurred while querying the RAG system."

    finally:
        # Ensure resources are released
        if db:
            del db  # Explicitly delete the database instance


# Script Execution Utilities
def run_script(script_name, arguments = "", cwd=None):
    """Run a Python script as a subprocess."""
    current_python = sys.executable
    if arguments == "":
        subprocess.run([current_python, script_name], cwd=cwd, check=True)
    else:
        subprocess.run([current_python, script_name, arguments], cwd=cwd, check=True)
        

# Save Results
import os
from datetime import datetime
from docx import Document

def save_questions_to_docx(questions, download_folder):
    """Save generated questions to a .docx file."""
    # Ensure the download folder exists
    os.makedirs(download_folder, exist_ok=True)

    # Create the .docx file with a timestamped filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"created_quiz_{timestamp}.docx"
    filepath = os.path.join(download_folder, filename)
    
    # Create and save the .docx file
    doc = Document()
    doc.add_heading('Generated Quiz', level=1)
    for question_type, question in questions:
        doc.add_paragraph(f"{question_type}: {question}")
    doc.save(filepath)
    
    return filename
