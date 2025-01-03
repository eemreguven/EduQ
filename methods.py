import os
import shutil
import subprocess
import re
import sys
from PyPDF2 import PdfReader
from docx import Document
from datetime import datetime
from flask import jsonify, url_for
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import (
    TranscriptsDisabled, NoTranscriptFound, VideoUnavailable
)
from youtube_transcript_api.formatters import TextFormatter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from werkzeug.utils import secure_filename
from constants import (
    PROMPT_TEMPLATE_FOR_QUESTIONS,
    PROMPT_TEMPLATE_FOR_PROMPTS, 
    ALLOWED_EXTENSIONS, 
    CHROMA_FOLDER_PATH,
    UPLOAD_FOLDER_PATH, 
    DOWNLOAD_FOLDER_PATH,
    question_types, 
    example_prompt_templates
)
from langchain_chroma import Chroma
from langchain.prompts import ChatPromptTemplate
from langchain_ollama import OllamaLLM
from progress import progress_data

# Add the 'rag-system' directory to sys.path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), 'rag-system')))
from get_embeddings import get_embeddings_function

#--------------------------------------------------------------------------------------------#

def handle_file_upload(file):
    """
    Handle the upload and preprocessing of a file.

    Args:
        file (FileStorage): The uploaded file object.

    Returns:
        Response: A JSON response indicating success or failure, with appropriate status codes.
    """
    if not file:
        return jsonify({'error': 'No file provided.'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': 'Unsupported file type.'}), 415

    try:
        # Secure the filename and save the file
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER_PATH, filename)
        os.makedirs(UPLOAD_FOLDER_PATH, exist_ok=True)
        file.save(file_path)

        # Preprocess the uploaded file
        preprocess_result = preprocess_file(file_path)
        if 'error' in preprocess_result:
            return jsonify(preprocess_result), 500

        # Return success response
        return jsonify({
            'success': True,
            'processing_message': preprocess_result['message'],
            'redirect_url': url_for('questions')
        }), 200

    except Exception as e:
        print(f"Exception in handle_file_upload: {e}")
        return jsonify({'error': f"An unexpected error occurred: {str(e)}"}), 500


def allowed_file(filename):
    """Check if the file type is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

#--------------------------------------------------------------------------------------------#

def handle_youtube_upload(youtube_url):
    """
    Handle the processing of a YouTube URL.

    Args:
        youtube_url (str): The YouTube URL provided by the user.

    Returns:
        Response: A JSON response indicating success or failure, with appropriate status codes.
    """
    if not youtube_url:
        return jsonify({'error': 'YouTube URL is required.'}), 400  # 400: Bad Request

    try:
        preprocess_result = preprocess_youtube(youtube_url)
        if preprocess_result.get('error'):
            return jsonify(preprocess_result), 400  # 400: Bad Request

        return jsonify({
            'success': True,
            'processing_message': preprocess_result['message'],
            'redirect_url': url_for('questions')
        }), 200  # 200: OK

    except Exception as e:
        return jsonify({'error': f"An unexpected error occurred: {str(e)}"}), 500  # 500: Internal Server Error


def preprocess_youtube(video_url):
    """
    Fetch and save a YouTube transcript.

    Args:
        video_url (str): The YouTube video URL.

    Returns:
        dict: A dictionary containing a success message or an error message.
    """
    try:
        transcript = fetch_youtube_transcript(video_url)
        if not transcript:
            return {'error': "No transcript available for this video."}

        video_title = "YouTube_Transcript"
        pdf_result = save_transcript_to_pdf(transcript, video_title)

        if pdf_result.get('error'):
            return pdf_result  # Return the error from saving the PDF

        return {'message': pdf_result['message']}

    except ValueError as e:
        return {'error': str(e)}  # Invalid YouTube URL

    except Exception as e:
        return {'error': f"An unexpected error occurred during preprocessing: {str(e)}"}


def fetch_youtube_transcript(video_url, language="en"):
    """
    Fetch the transcript for a YouTube video.

    Args:
        video_url (str): The YouTube video URL.
        language (str): The desired language for the transcript (default: "en").

    Returns:
        str: The cleaned transcript, or None if not available.
    """
    try:
        video_id = extract_video_id(video_url)
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
        transcript = transcript_list.find_transcript([language])
        formatter = TextFormatter()
        raw_transcript = formatter.format_transcript(transcript.fetch())

        # Clean up the transcript
        return clean_transcript(raw_transcript)
    except NoTranscriptFound:
        return None  # No transcript available
    except TranscriptsDisabled:
        return None  # Transcripts disabled for the video
    except VideoUnavailable:
        return None  # Video is unavailable
    except Exception as e:
        print(f"Error fetching transcript: {e}")
        return None


def save_transcript_to_pdf(transcript, video_title):
    """
    Save a cleaned transcript to a PDF file.

    Args:
        transcript (str): The cleaned transcript text.
        video_title (str): The title of the video.

    Returns:
        dict: A dictionary containing a success message or an error message.
    """
    pdf_file = None  # Initialize to ensure clean handling
    try:
        # Ensure the upload folder exists
        os.makedirs(UPLOAD_FOLDER_PATH, exist_ok=True)
        pdf_path = os.path.join(UPLOAD_FOLDER_PATH, f"{video_title.replace(' ', '_')}.pdf")

        # Safely create the PDF
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, 750, "Transcription")
        c.setFont("Helvetica", 12)
        y_position = 730
        line_height = 14

        # Write transcript line by line
        for line in transcript.split("\n"):
            if y_position <= 40:  # New page when content exceeds bounds
                c.showPage()
                c.setFont("Helvetica", 12)
                y_position = 750
            c.drawString(50, y_position, line)
            y_position -= line_height

        # Finalize and save PDF
        c.save()

        # Explicitly check if the file exists and is closed
        if os.path.isfile(pdf_path):
            pdf_file = open(pdf_path, "rb")
            pdf_file.close()

        return {'message': f"Transcription saved to {pdf_path}."}

    except Exception as e:
        return {'error': f"Error saving transcript to PDF: {str(e)}"}

    finally:
        # Ensure cleanup in case of an unexpected failure
        if pdf_file and not pdf_file.closed:
            pdf_file.close()


def extract_video_id(video_url):
    """
    Extract the video ID from a YouTube URL.

    Args:
        video_url (str): The YouTube URL.

    Returns:
        str: The extracted video ID.

    Raises:
        ValueError: If the URL is invalid.
    """
    if "v=" in video_url:
        return video_url.split("v=")[-1].split("&")[0]
    elif "youtu.be/" in video_url:
        return video_url.split("youtu.be/")[-1].split("?")[0]
    else:
        raise ValueError("Invalid YouTube URL")


def clean_transcript(transcript, max_line_length=80):
    """
    Clean the YouTube transcript by removing unnecessary newlines
    and formatting sentences to avoid line overflow.

    Args:
        transcript (str): The raw transcript text.
        max_line_length (int): The maximum number of characters per line.

    Returns:
        str: A cleaned transcript with properly formatted lines.
    """
    # Remove excessive newlines and combine broken lines
    cleaned_lines = []
    current_sentence = ""

    for line in transcript.splitlines():
        line = line.strip()  # Remove leading/trailing spaces
        if line:  # If line is not empty
            if current_sentence:
                current_sentence += " " + line
            else:
                current_sentence = line

            # If the current line ends with punctuation, finalize the sentence
            if re.search(r'[.!?]$', line):
                cleaned_lines.extend(wrap_line(current_sentence, max_line_length))
                current_sentence = ""

    # Append the last sentence if any
    if current_sentence:
        cleaned_lines.extend(wrap_line(current_sentence, max_line_length))

    return "\n".join(cleaned_lines)


def wrap_line(sentence, max_length):
    """
    Wrap a single sentence into multiple lines without splitting words.

    Args:
        sentence (str): The sentence to wrap.
        max_length (int): The maximum number of characters per line.

    Returns:
        list: A list of lines, each with a maximum length of `max_length`.
    """
    words = sentence.split()
    lines = []
    current_line = []

    for word in words:
        # Check if adding the next word exceeds the max_length
        if sum(len(w) for w in current_line) + len(current_line) + len(word) > max_length:
            lines.append(" ".join(current_line))
            current_line = [word]  # Start a new line with the current word
        else:
            current_line.append(word)

    # Add the last line if any
    if current_line:
        lines.append(" ".join(current_line))

    return lines

#--------------------------------------------------------------------------------------------#

def preprocess_file(file_path):
    """
    Preprocess files by verifying their existence.

    Args:
        file_path (str): The path to the file to be processed.

    Returns:
        dict: A dictionary with 'message' for success or 'error' for failure.
    """
    try:
        # Check if the file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Perform minimal preprocessing or verification here if needed
        return {'message': f"File '{os.path.basename(file_path)}' uploaded successfully."}

    except FileNotFoundError as e:
        return {'error': str(e)}

    except Exception as e:
        return {'error': f"An unexpected error occurred: {str(e)}"}

#--------------------------------------------------------------------------------------------#

def validate_questions(form_data):
    """
    Validate user input for question generation.

    Args:
        form_data (ImmutableMultiDict): The form data submitted by the user.

    Returns:
        tuple: A tuple containing validated question data and a list of errors.
    """
    errors = []
    question_data = []

    for i, question_type in enumerate(question_types, start=1):
        try:
            # Parse input values
            easy_count = int(form_data.get(f'easy_{i}', 0))
            medium_count = int(form_data.get(f'medium_{i}', 0))
            difficult_count = int(form_data.get(f'difficult_{i}', 0))

            # Validate the counts
            for difficulty, count in [("easy", easy_count), ("medium", medium_count), ("difficult", difficult_count)]:
                if not (0 <= count <= 10):
                    errors.append(f"{difficulty.capitalize()} count for '{question_type}' must be between 0 and 10.")

            question_data.append({
                'question_type': question_type,
                'easy': easy_count,
                'medium': medium_count,
                'difficult': difficult_count
            })

        except ValueError:
            errors.append(f"Invalid input for '{question_type}'. Please enter numeric values.")

    return question_data, errors


def generate_questions(question_data):
    """
    Generate questions using the RAG system and save them to a .docx file.

    Args:
        question_data (list): List of dictionaries containing question configurations.

    Returns:
        str: The filename of the saved .docx file.
    """
    try:
        progress_data["status"] = "Updating the resource database..."
        print("Populating the database...")
        run_script('populate_database.py', cwd='rag-system')

        progress_data["status"] = f"Generating summary of documents to create dynamic prompts."
        print(progress_data["status"])
        summary = get_summary_of_all_documents()

        generated_questions = []
        total_questions = sum(item[difficulty] for item in question_data for difficulty in ['easy', 'medium', 'difficult'])
        current_question = 0
                
        # Generate questions with progress updates
        for item in question_data:
            question_type = item['question_type']
            for difficulty in ['easy', 'medium', 'difficult']:
                count = item[difficulty]
                if count > 0:
                    for _ in range(count):
                        current_question += 1
                        progress_data["status"] = f"Generating question {current_question}/{total_questions} ({question_type} - {difficulty.capitalize()})"
                        print(progress_data["status"])
                        
                        # Generate dynamic query and fetch the question response
                        dynamic_query =  create_dynamic_query(summary, question_type, difficulty)
                        question_response = query_rag(dynamic_query)
                        generated_questions.append((f"{question_type} {difficulty.capitalize()}", question_response))

        progress_data["status"] = "Saving generated questions..."
        filename = save_questions_to_docx(generated_questions)
        progress_data["status"] = "Questions generated successfully."
        return filename

    except Exception as e:
        progress_data["status"] = "Error during question generation."
        print(f"Error in generate_questions: {str(e)}")
        raise RuntimeError("Failed to generate questions.")


def query_rag(query_text, chroma_path=CHROMA_FOLDER_PATH, prompt_template=PROMPT_TEMPLATE_FOR_QUESTIONS):
    """
    Query the RAG system for context-based answers.

    Args:
        query_text (str): The query text.
        chroma_path (str): The path to the RAG system database.
        prompt_template (str): Template for the query.

    Returns:
        str: The generated question or an error message.
    """
    try:
        print(f"Querying... Text: {query_text}")

        embedding_function = get_embeddings_function()
        db = Chroma(persist_directory=chroma_path, embedding_function=embedding_function)

        results = db.similarity_search_with_score(query_text, k=5)

        if not results:
            return "No relevant context found."

        context_text = "\n\n---\n\n".join([doc.page_content for doc, _ in results])
        prompt = ChatPromptTemplate.from_template(prompt_template).format(
            context=context_text, question=query_text
        )

        model = OllamaLLM(model="llama3.2:3b")
        return model.invoke(prompt)

    except Exception as e:
        progress_data["status"] = f"Error querying RAG system: {str(e)}"
        print(progress_data["status"])
        return "An error occurred while querying the RAG system."

def create_dynamic_query(summary, question_type, difficulty, chroma_path=CHROMA_FOLDER_PATH, prompt_template=PROMPT_TEMPLATE_FOR_PROMPTS):
    """
    Create a dynamic query using the document summary, question type, and difficulty.

    Args:
        summary (str): A summary of the document context.
        question_type (str): The type of question (e.g., "True/False", "Multiple Choice").
        difficulty (str): The difficulty level (e.g., "easy", "medium", "difficult").
        chroma_path (str): Path to the Chroma database.
        prompt_template (str): Template to guide LLM in generating the query.

    Returns:
        str: A formatted query for generating questions.
    """
    try:        
        # Get the example prompt for the question type and difficulty
        example_prompt = example_prompt_templates[question_type][difficulty]
        question_base = f"{question_type}-{difficulty}"
        
        # Retrieve relevant contexts from the database
        embedding_function = get_embeddings_function()
        db = Chroma(persist_directory=chroma_path, embedding_function=embedding_function)
        
        num_docs_to_include = db._collection.count()
        results = db.similarity_search_with_score(summary, k=num_docs_to_include)
        
        if not results:
            return "No relevant context found to create the prompt template."
        
        # Combine the retrieved contexts
        context_text = "\n\n---\n\n".join([doc.page_content for doc, _ in results])
        
        # Format the dynamic prompt template
        formatted_prompt = ChatPromptTemplate.from_template(prompt_template).format(
            context=context_text,
            question_base=question_base,
            example_prompt=example_prompt
        )
        
        # Query the LLM for a dynamic prompt
        model = OllamaLLM(model="llama3.2:3b")
        dynamic_prompt = model.invoke(formatted_prompt).strip()
        
        # Format the query for question generation
        question_format = example_prompt_templates[question_type]["question_format"]
        query_text = f"{question_format}\n\n{dynamic_prompt}"
        
        return query_text
    
    except Exception as e:
        print(f"Error generating dynamic prompt template: {e}")
        return "Error generating prompt template."


def get_summary_of_all_documents(max_docs=None):
    """
    Summarize the content of all documents in the Chroma database.

    Args:
        max_docs (int, optional): Limit the number of documents to include in the summary. Defaults to None.

    Returns:
        str: A summary of all documents in the database.
    """
    try:
        # Connect to the Chroma database
        db = Chroma(persist_directory=CHROMA_FOLDER_PATH, embedding_function=get_embeddings_function())

        # Fetch all documents
        all_documents = db.get(include=["documents"])
        total_docs = len(all_documents["documents"])

        # Limit the number of documents if max_docs is provided
        if max_docs is not None and max_docs < total_docs:
            documents = all_documents["documents"][:max_docs]
        else:
            documents = all_documents["documents"]

        # Combine content from all documents
        combined_content = "\n\n---\n\n".join(documents)

        # Create a prompt for summarization
        summary_prompt = f"Summarize the following content in 5 sentences:\n\n{combined_content}"

        # Use the LLM to generate the summary
        model = OllamaLLM(model="llama3.2:3b")
        summary = model.invoke(summary_prompt).strip()
        
        return summary

    except Exception as e:
        print(f"Error generating summary of documents: {e}")
        return "An error occurred while summarizing the documents."


def save_questions_to_docx(questions_and_answers):
    """
    Save processed questions and answers to a .docx file.

    Args:
        questions_and_answers (list): A list of tuples containing question type, question text, and answer.

    Returns:
        str: The filename of the saved .docx file.
    """
    try:
        os.makedirs(DOWNLOAD_FOLDER_PATH, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"quiz_{timestamp}.docx"
        filepath = os.path.join(DOWNLOAD_FOLDER_PATH, filename)

        # Write to .docx
        doc = Document()
        doc.add_heading('Generated Questions', level=2)

        for idx, (question_type, questions_and_answer) in enumerate(questions_and_answers, 1):
            doc.add_paragraph(f"{idx}-)({question_type}) {questions_and_answer}")
        doc.save(filepath)
        return filename
    except Exception as e:
        print(f"Error saving questions to DOCX: {str(e)}")
        raise RuntimeError("Failed to save questions to file.")


def reset_database(chroma_path=CHROMA_FOLDER_PATH):
    """
    Delete all vectors (documents, embeddings, etc.) from the Chroma database.

    Args:
        chroma_path (str): The path to the Chroma database.
    """
    try:
        # Connect to the Chroma database
        db = Chroma(persist_directory=chroma_path, embedding_function=get_embeddings_function())


        db.reset_collection()
    except Exception as e:
        print(f"Error deleting all entries from Chroma: {e}")


def count_characters(file_path, file_type):
    """Count the number of characters in a file based on its type."""
    try:
        if file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return len(f.read())
        elif file_type == 'pdf':
            reader = PdfReader(file_path)
            return sum(len(page.extract_text()) for page in reader.pages)
        elif file_type in {'doc', 'docx'}:
            doc = Document(file_path)
            return sum(len(paragraph.text) for paragraph in doc.paragraphs)
        return 0
    except Exception as e:
        return f"Error counting characters: {e}"

#--------------------------------------------------------------------------------------------#

def run_script(script_name, arguments="", cwd=None):
    """
    Run a Python script as a subprocess.

    Args:
        script_name (str): The name of the script to run.
        arguments (str, optional): Additional arguments for the script.
        cwd (str, optional): Working directory for the script.

    Raises:
        RuntimeError: If the script execution fails.
    """
    try:
        current_python = sys.executable
        cmd = [current_python, script_name] + arguments.split()
        subprocess.run(cmd, cwd=cwd, check=True)
    except subprocess.CalledProcessError as e:
        print(f"Error running script '{script_name}': {str(e)}")
        raise RuntimeError(f"Script execution failed: {script_name}")

#--------------------------------------------------------------------------------------------#

def clear_folder(folder_path):
    """
    Delete all files and directories in the specified folder.

    Args:
        folder_path (str): The path to the folder to clear.
    """
    try:
        if os.path.exists(folder_path):
            for file_name in os.listdir(folder_path):
                file_path = os.path.join(folder_path, file_name)
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)  # Remove the file or link
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)  # Remove the directory
            print(f"Cleared all files in: {folder_path}")
        else:
            print(f"Folder does not exist: {folder_path}")
    except Exception as e:
        print(f"Error clearing folder '{folder_path}': {str(e)}")
